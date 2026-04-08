# documents/ai_service.py
"""
Service AI unifie pour AltiusOne — 100% local.

- OCR: Tesseract local (images, PDF scannés, multilingue FR/DE/EN/IT)
- Embeddings: sentence-transformers local (768D) via core.ai.embeddings
- Chat/Classification/Extraction: Ollama local via core.ai.chat
- Résumé et Q&A sur documents

Aucune dépendance externe. Aucune donnée ne quitte l'instance.
"""
import logging
import hashlib
from typing import Optional, Dict, Any, List
from dataclasses import dataclass
from django.conf import settings

# Import des schemas personnalises
from documents.schemas import (
    DOCUMENT_TYPE_SCHEMAS,
    get_schema_for_document_type,
    get_available_document_types,
)

logger = logging.getLogger(__name__)


@dataclass
class OCRResult:
    """Resultat d'extraction OCR."""
    text: str
    confidence: float = 0.0
    pages: int = 1
    language: str = 'fr'
    processing_time: float = 0.0


@dataclass
class ClassificationResult:
    """Resultat de classification de document."""
    type_document: str
    confidence: float
    categories: List[str]
    tags: List[str]


@dataclass
class ExtractionResult:
    """Resultat d'extraction de metadonnees."""
    data: Dict[str, Any]
    confidence: float = 0.0


class AIServiceError(Exception):
    """Exception pour les erreurs du service AI."""
    pass


class AltiusAIService:
    """
    Service unifie pour toutes les operations AI via le SDK AltiusOne.

    Utilise l'API AltiusOne AI pour:
    - OCR intelligent
    - Generation d'embeddings 768D
    - Extraction de donnees structurees
    - Chat IA contextuel
    """

    # Mapping types de documents vers categories
    TYPE_TO_CATEGORY = {
        # Comptabilite
        'FACTURE': 'Comptabilite',
        'FACTURE_ACHAT': 'Comptabilite',
        'FACTURE_VENTE': 'Comptabilite',
        'DEVIS': 'Commercial',
        'OFFRE': 'Commercial',
        'BON_COMMANDE': 'Commercial',
        'BON_LIVRAISON': 'Logistique',

        # Banque
        'RELEVE_BANQUE': 'Banque',
        'EXTRAIT_BANCAIRE': 'Banque',

        # Contrats
        'CONTRAT': 'Juridique',
        'CONTRAT_TRAVAIL': 'RH',
        'CONTRAT_BAIL': 'Juridique',

        # RH / Salaires
        'FICHE_SALAIRE': 'RH',
        'CERTIFICAT_SALAIRE': 'RH',

        # Fiscal
        'DECLARATION_TVA': 'Fiscal',

        # Administratif
        'ATTESTATION': 'Administratif',
        'CORRESPONDANCE': 'Correspondance',
        'COURRIER': 'Correspondance',
        'EMAIL': 'Correspondance',

        # Divers
        'AUTRE': 'Divers',
    }

    def __init__(self):
        self._embedding_service = None
        self._chat_service = None

    @property
    def embedding_svc(self):
        if self._embedding_service is None:
            from core.ai.embeddings import embedding_service
            self._embedding_service = embedding_service
        return self._embedding_service

    @property
    def chat_svc(self):
        if self._chat_service is None:
            from core.ai.chat import chat_service
            self._chat_service = chat_service
        return self._chat_service

    @property
    def enabled(self) -> bool:
        """Le service local est toujours disponible."""
        return True

    # _make_request supprimé — plus d'API externe

    # =========================================================================
    # OCR - EXTRACTION DE TEXTE
    # =========================================================================

    # Mapping langue -> code Tesseract
    TESSERACT_LANGS = {
        'fr': 'fra', 'de': 'deu', 'it': 'ita', 'en': 'eng',
        'auto': 'fra+deu+eng+ita',
    }

    def ocr(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        file_url: Optional[str] = None,
        filename: Optional[str] = None,
        mime_type: Optional[str] = None,
        language: str = 'auto'
    ) -> OCRResult:
        """
        Extrait le texte d'un document (image, PDF ou DOCX).

        Utilise Tesseract OCR localement (rapide, CPU).
        DOCX: extraction directe du texte (pas d'OCR).
        PDF: extraction du texte embarque, puis OCR sur les pages images.
        Images: OCR Tesseract.

        Args:
            file_path: Chemin vers le fichier local
            file_content: Contenu du fichier en bytes
            file_url: URL du fichier (telecharge puis traite)
            filename: Nom du fichier (pour determiner le mime_type)
            mime_type: Type MIME du fichier (prioritaire sur la detection)
            language: Langue du document (auto, fr, de, en, it)

        Returns:
            OCRResult avec le texte extrait et metadonnees
        """
        import time as _time
        start = _time.time()

        try:
            if file_path:
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                if not filename:
                    filename = file_path.split('/')[-1]

            if file_url and not file_content:
                import requests as _req
                resp = _req.get(file_url, timeout=30)
                resp.raise_for_status()
                file_content = resp.content
                if not filename:
                    filename = file_url.split('/')[-1].split('?')[0] or 'document'

            if not file_content:
                raise AIServiceError("Aucune source de fichier fournie")

            import mimetypes
            if not filename:
                filename = 'document'
            if not mime_type:
                mime_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'

            # DOCX — extraction directe
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or filename.lower().endswith('.docx'):
                return self._extract_text_from_docx(file_content, filename)

            # PDF — texte embarque + OCR fallback sur pages images
            if 'pdf' in mime_type.lower() or filename.lower().endswith('.pdf'):
                return self._ocr_pdf(file_content, language, start)

            # Images — OCR Tesseract
            return self._ocr_image(file_content, language, start)

        except AIServiceError:
            raise
        except Exception as e:
            logger.error(f"Erreur OCR: {e}")
            raise AIServiceError(f"Erreur lors de l'extraction OCR: {str(e)}")

    def _ocr_image(self, image_bytes: bytes, language: str, start_time: float) -> OCRResult:
        """OCR d'une image avec Tesseract — scores de confiance réels."""
        import time as _time
        import pytesseract
        from PIL import Image
        import io

        lang = self.TESSERACT_LANGS.get(language, 'fra+deu+eng+ita')
        img = Image.open(io.BytesIO(image_bytes))

        text = pytesseract.image_to_string(img, lang=lang)

        # Score de confiance réel depuis Tesseract (mot par mot)
        confidence = self._get_tesseract_confidence(img, lang)

        processing_ms = (_time.time() - start_time) * 1000
        detected_lang = language if language != 'auto' else 'fr'

        return OCRResult(
            text=text.strip(),
            confidence=confidence,
            pages=1,
            language=detected_lang,
            processing_time=processing_ms
        )

    @staticmethod
    def _get_tesseract_confidence(img, lang: str) -> float:
        """Calcule le score de confiance moyen réel depuis Tesseract."""
        import pytesseract
        try:
            data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
            confidences = [
                int(c) for c in data.get('conf', [])
                if str(c).lstrip('-').isdigit() and int(c) > 0
            ]
            if confidences:
                return round(sum(confidences) / len(confidences), 1)
        except Exception:
            pass
        return 0.0

    def _ocr_pdf(self, pdf_bytes: bytes, language: str, start_time: float) -> OCRResult:
        """OCR d'un PDF: texte embarque d'abord, puis OCR sur pages images."""
        import time as _time
        import io

        # 1. Essayer d'extraire le texte embarque (PDF textuel)
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(pdf_bytes))
            num_pages = len(reader.pages)
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())

            # Si on a du texte pour la majorite des pages, c'est un PDF textuel
            if len(text_parts) >= num_pages * 0.5 and any(len(t) > 20 for t in text_parts):
                processing_ms = (_time.time() - start_time) * 1000
                return OCRResult(
                    text='\n\n'.join(text_parts),
                    confidence=95.0,
                    pages=num_pages,
                    language=language if language != 'auto' else 'fr',
                    processing_time=processing_ms
                )
        except Exception as e:
            logger.warning(f"Extraction texte PDF echouee, fallback OCR: {e}")
            num_pages = 0

        # 2. PDF image — convertir en images et OCR chaque page
        import pytesseract
        from pdf2image import convert_from_bytes

        lang = self.TESSERACT_LANGS.get(language, 'fra+deu+eng+ita')
        images = convert_from_bytes(pdf_bytes, dpi=300)
        num_pages = len(images)
        text_parts = []
        page_confidences = []

        for img in images:
            page_text = pytesseract.image_to_string(img, lang=lang)
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
                page_confidences.append(self._get_tesseract_confidence(img, lang))

        processing_ms = (_time.time() - start_time) * 1000

        # Confiance moyenne sur toutes les pages
        avg_confidence = (
            round(sum(page_confidences) / len(page_confidences), 1)
            if page_confidences else 0.0
        )

        return OCRResult(
            text='\n\n'.join(text_parts),
            confidence=avg_confidence,
            pages=num_pages,
            language=language if language != 'auto' else 'fr',
            processing_time=processing_ms
        )

    def _extract_text_from_docx(self, file_content: bytes, filename: str) -> OCRResult:
        """
        Extrait le texte d'un fichier DOCX localement (sans OCR).

        Args:
            file_content: Contenu du fichier DOCX en bytes
            filename: Nom du fichier

        Returns:
            OCRResult avec le texte extrait
        """
        import io
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise AIServiceError("python-docx non installe. Installez-le avec: pip install python-docx")

        try:
            # Charger le document DOCX depuis les bytes
            doc = DocxDocument(io.BytesIO(file_content))

            # Extraire le texte de tous les paragraphes
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extraire aussi le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))

            text = '\n'.join(paragraphs)

            return OCRResult(
                text=text,
                confidence=100.0,  # Extraction directe, pas d'OCR
                pages=1,
                language='fr',
                processing_time=0.0
            )

        except Exception as e:
            logger.error(f"Erreur extraction DOCX {filename}: {e}")
            raise AIServiceError(f"Erreur lors de l'extraction du texte DOCX: {str(e)}")

    # =========================================================================
    # EMBEDDINGS - VECTORISATION (local sentence-transformers)
    # =========================================================================

    def embed(self, text: str) -> List[float]:
        """Genere un embedding 768D via sentence-transformers local."""
        if not text or not text.strip():
            raise AIServiceError("Texte vide fourni pour l'embedding")
        result = self.embedding_svc.generate_embedding(text)
        if result is None:
            raise AIServiceError("Echec generation embedding")
        return result

    def embed_batch(self, texts: List[str]) -> List[Optional[List[float]]]:
        """Genere des embeddings pour plusieurs textes."""
        return self.embedding_svc.generate_embeddings_batch(texts)

    @property
    def embedding_dimensions(self) -> int:
        return self.embedding_svc.dimensions

    # =========================================================================
    # CLASSIFICATION DE DOCUMENTS
    # =========================================================================

    def classify_document(
        self,
        text: str,
        filename: Optional[str] = None
    ) -> ClassificationResult:
        """
        Classifie un document par similarite vectorielle (pas de LLM).

        Compare l'embedding du document avec des embeddings de reference
        pour chaque type. ~50ms, deterministe.

        La metrique de distance est configurable via le classifier
        (cosine par defaut, extensible a l2, jaccard, etc.)

        Args:
            text: Texte du document (extrait par OCR)
            filename: Nom du fichier (optionnel, ajoute au texte)

        Returns:
            ClassificationResult avec type, confiance et tags
        """
        from core.ai.classifier import document_classifier

        try:
            # Preparer le texte (inclure le nom de fichier pour aider)
            classify_text = text[:8000]
            if filename:
                classify_text = f"Fichier: {filename}\n\n{classify_text}"

            # Classification par similarite vectorielle
            results = document_classifier.classify(classify_text, top_k=3)

            if not results:
                return ClassificationResult(
                    type_document='AUTRE', confidence=0.0,
                    categories=['Divers'], tags=[]
                )

            best_type, best_score = results[0]
            category = self.TYPE_TO_CATEGORY.get(best_type, 'Divers')

            # Tags = les 2-3 meilleurs types si proches en score
            tags = []
            for doc_type, score in results[1:]:
                if score > best_score * 0.8:  # dans 80% du meilleur score
                    tags.append(doc_type.lower())

            return ClassificationResult(
                type_document=best_type,
                confidence=round(best_score, 3),
                categories=[category],
                tags=tags
            )

        except Exception as e:
            logger.error(f"Erreur classification document: {e}")
            return ClassificationResult(
                type_document='AUTRE', confidence=0.0,
                categories=['Divers'], tags=[]
            )

    # =========================================================================
    # EXTRACTION DE METADONNEES STRUCTUREES
    # =========================================================================

    def extract_metadata(
        self,
        text: str,
        type_document: str,
        custom_schema: Optional[Dict] = None,
        source_language: str = 'auto',
        output_language: str = 'fr'
    ) -> ExtractionResult:
        """
        Extrait les metadonnees structurees d'un document.

        Utilise les schemas personnalises definis dans documents/schemas.py
        pour extraire des informations structurees (adresses, montants,
        numeros TVA, AVS, noms, etc.)

        Args:
            text: Texte du document
            type_document: Type de document (FACTURE_ACHAT, CONTRAT_TRAVAIL, etc.)
            custom_schema: Schema personnalise (prioritaire sur le schema predefini)
            source_language: Langue du document source (auto, fr, de, it, en)
            output_language: Langue de sortie pour les extractions (fr par defaut)

        Returns:
            ExtractionResult avec les donnees extraites
        """
        schema = custom_schema or get_schema_for_document_type(type_document)

        import json as _json
        system_prompt = f"""Tu es un assistant d'extraction de metadonnees pour documents professionnels suisses.
Extrais les informations structurees du document selon ce schema JSON:
{_json.dumps(schema, indent=2, ensure_ascii=False) if schema else 'Extrais toutes les informations pertinentes.'}

Reponds UNIQUEMENT en JSON valide. Langue de sortie: {output_language}."""

        user_prompt = f"Document de type {type_document}:\n\n{text[:8000]}"

        try:
            chat_response = self.chat(message=user_prompt, system=system_prompt, temperature=0.1)
            response_text = chat_response.get('response', '')

            try:
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                if json_start >= 0 and json_end > json_start:
                    data = _json.loads(response_text[json_start:json_end])
                else:
                    data = {}
            except _json.JSONDecodeError:
                data = {}

            return ExtractionResult(data=data, confidence=0.8 if data else 0.0)

        except Exception as e:
            logger.error(f"Erreur extraction metadonnees: {e}")
            return ExtractionResult(data={}, confidence=0.0)

    # =========================================================================
    # CHAT IA
    # =========================================================================

    def chat(
        self,
        message: str = "",
        system: Optional[str] = None,
        system_prompt: Optional[str] = None,
        context: Optional[str] = None,
        history: Optional[List[Dict[str, str]]] = None,
        temperature: float = 0.7,
        max_tokens: Optional[int] = None,
        tools: Optional[List[dict]] = None,
        messages_override: Optional[List[Dict[str, str]]] = None,
    ) -> Dict[str, Any]:
        """Interroge le LLM local via Ollama. Même interface qu'avant."""
        try:
            return self.chat_svc.chat(
                message=message,
                system=system,
                system_prompt=system_prompt,
                context=context,
                history=history,
                temperature=temperature,
                max_tokens=max_tokens,
                tools=tools,
                messages_override=messages_override,
            )
        except Exception as e:
            logger.error(f"Erreur chat IA: {e}")
            raise AIServiceError(str(e))

    def chat_stream(
        self,
        message: str = "",
        system: Optional[str] = None,
        system_prompt: Optional[str] = None,
        context: Optional[str] = None,
        history: Optional[List[Dict[str, str]]] = None,
        temperature: float = 0.7,
        max_tokens: Optional[int] = None,
        tools: Optional[List[dict]] = None,
        messages_override: Optional[List[Dict[str, str]]] = None,
    ):
        """Stream la réponse token par token via Ollama local."""
        yield from self.chat_svc.chat_stream(
            message=message,
            system=system,
            system_prompt=system_prompt,
            context=context,
            history=history,
            temperature=temperature,
            max_tokens=max_tokens,
            tools=tools,
            messages_override=messages_override,
        )

    # =========================================================================
    # RESUME ET Q&A SUR DOCUMENTS
    # =========================================================================

    def summarize_document(
        self,
        text: str,
        max_length: str = 'medium',
        language: str = 'fr'
    ) -> Dict[str, Any]:
        """
        Genere un resume d'un document (utile pour documents longs).

        Args:
            text: Texte complet du document
            max_length: Longueur du resume ('short', 'medium', 'long')
            language: Langue du resume (fr, de, en, it)

        Returns:
            Dict avec 'summary', 'key_points', 'entities'
        """
        length_instructions = {
            'short': '3-5 phrases maximum',
            'medium': '1-2 paragraphes',
            'long': '3-4 paragraphes detailles'
        }

        system_prompt = f"""Tu es un assistant specialise dans l'analyse de documents professionnels suisses.
Genere un resume clair et structure en {language.upper()}.
Longueur attendue: {length_instructions.get(max_length, '1-2 paragraphes')}

Format de reponse JSON:
{{
    "summary": "Resume du document...",
    "key_points": ["Point cle 1", "Point cle 2", ...],
    "entities": {{
        "personnes": ["Nom 1", "Nom 2"],
        "organisations": ["Entreprise 1"],
        "montants": ["CHF 1'000.00"],
        "dates": ["15.01.2024"]
    }},
    "type_document_suggere": "FACTURE_ACHAT"
}}"""

        # Limiter le texte pour eviter les timeouts (4000 chars max)
        text_truncated = text[:4000] if len(text) > 4000 else text

        user_prompt = f"""Resume ce document:

{text_truncated}"""

        try:
            chat_response = self.chat(message=user_prompt, system=system_prompt, temperature=0.3, max_tokens=1000)
            response_text = chat_response.get('response', '') if isinstance(chat_response, dict) else str(chat_response)

            # Parser la reponse JSON
            import json
            try:
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                if json_start >= 0 and json_end > json_start:
                    result = json.loads(response_text[json_start:json_end])
                else:
                    result = {'summary': response_text, 'key_points': [], 'entities': {}}
            except json.JSONDecodeError:
                result = {'summary': response_text, 'key_points': [], 'entities': {}}

            return result

        except Exception as e:
            logger.error(f"Erreur resume document: {e}")
            return {'summary': '', 'key_points': [], 'entities': {}, 'error': str(e)}

    def ask_document(
        self,
        text: str,
        question: str,
        history: Optional[List[Dict[str, str]]] = None,
        language: str = 'fr'
    ) -> Dict[str, Any]:
        """
        Pose une question sur un document specifique (Q&A contextuel).

        Permet aux utilisateurs de poser des questions sur le contenu
        d'un document et d'obtenir des reponses precises basees sur le texte.

        Args:
            text: Texte du document (OCR)
            question: Question de l'utilisateur
            history: Historique de conversation pour le suivi
            language: Langue de reponse

        Returns:
            Dict avec 'answer', 'confidence', 'sources'
        """
        system_prompt = f"""Tu es un assistant expert pour analyser des documents professionnels suisses.
Tu reponds aux questions en te basant UNIQUEMENT sur le contenu du document fourni.
Si l'information n'est pas dans le document, dis-le clairement.
Reponds en {language.upper()} de maniere precise et professionnelle.

Format de reponse JSON:
{{
    "answer": "Reponse detaillee...",
    "confidence": 0.95,
    "sources": ["Extrait pertinent 1 du document...", "Extrait pertinent 2..."],
    "found_in_document": true
}}"""

        # Limiter le texte pour eviter les timeouts (4000 chars max)
        text_truncated = text[:4000] if len(text) > 4000 else text

        user_prompt = f"""Document:
---
{text_truncated}
---

Question: {question}"""

        try:
            chat_response = self.chat(
                message=user_prompt,
                system=system_prompt,
                history=history,
                temperature=0.2,
                max_tokens=800  # Limiter pour eviter les timeouts
            )
            response_text = chat_response.get('response', '') if isinstance(chat_response, dict) else str(chat_response)

            # Parser la reponse JSON
            import json
            try:
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                if json_start >= 0 and json_end > json_start:
                    result = json.loads(response_text[json_start:json_end])
                else:
                    result = {'answer': response_text, 'confidence': 0.5, 'sources': [], 'found_in_document': True}
            except json.JSONDecodeError:
                result = {'answer': response_text, 'confidence': 0.5, 'sources': [], 'found_in_document': True}

            return result

        except Exception as e:
            logger.error(f"Erreur Q&A document: {e}")
            return {'answer': '', 'confidence': 0, 'sources': [], 'error': str(e)}

    # =========================================================================
    # TRAITEMENT COMPLET DE DOCUMENT
    # =========================================================================

    def process_document(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        filename: Optional[str] = None,
        mime_type: Optional[str] = None,
        generate_embedding: bool = True
    ) -> Dict[str, Any]:
        """
        Traitement complet d'un document:
        1. OCR (extraction texte)
        2. Classification
        3. Extraction metadonnees
        4. Generation embedding (optionnel)

        Args:
            file_path: Chemin vers le fichier
            file_content: Contenu en bytes
            filename: Nom du fichier
            mime_type: Type MIME
            generate_embedding: Generer l'embedding vectoriel

        Returns:
            Dict avec tous les resultats du traitement
        """
        result = {
            'ocr': None,
            'classification': None,
            'metadata': None,
            'embedding': None,
            'success': False,
            'errors': []
        }

        # 1. OCR
        try:
            ocr_result = self.ocr(
                file_path=file_path,
                file_content=file_content,
                filename=filename,
                mime_type=mime_type
            )
            result['ocr'] = {
                'text': ocr_result.text,
                'confidence': ocr_result.confidence,
                'pages': ocr_result.pages,
                'language': ocr_result.language
            }
        except AIServiceError as e:
            result['errors'].append(f"OCR: {str(e)}")
            return result

        text = ocr_result.text

        # 2. Classification
        try:
            classification = self.classify_document(text, filename)
            result['classification'] = {
                'type_document': classification.type_document,
                'confidence': classification.confidence,
                'categories': classification.categories,
                'tags': classification.tags
            }
        except Exception as e:
            result['errors'].append(f"Classification: {str(e)}")
            result['classification'] = {
                'type_document': 'AUTRE',
                'confidence': 0,
                'categories': ['Divers'],
                'tags': []
            }

        # 3. Extraction metadonnees
        try:
            type_doc = result['classification']['type_document']
            extraction = self.extract_metadata(text, type_doc)
            result['metadata'] = extraction.data
        except Exception as e:
            result['errors'].append(f"Extraction: {str(e)}")
            result['metadata'] = {}

        # 4. Embedding
        if generate_embedding:
            try:
                embedding = self.embed(text)
                result['embedding'] = embedding
            except Exception as e:
                result['errors'].append(f"Embedding: {str(e)}")

        result['success'] = len(result['errors']) == 0
        return result

    # =========================================================================
    # UTILITAIRES
    # =========================================================================

    def compute_text_hash(self, text: str) -> str:
        """Calcule le hash SHA-256 d'un texte."""
        return hashlib.sha256(text.encode()).hexdigest()

    def compute_similarity(self, embedding1: List[float], embedding2: List[float]) -> float:
        """Calcule la similarite cosinus entre deux embeddings."""
        return self.embedding_svc.compute_similarity(embedding1, embedding2)

    def health_check(self) -> Dict[str, Any]:
        """Verifie l'etat des services IA locaux."""
        status = {
            'enabled': True,
            'backend': 'local',
            'embedding_available': self.embedding_svc.is_available(),
            'embedding_model': self.embedding_svc.model_name,
            'connected': False,
            'error': None,
        }

        try:
            # Test embedding
            test = self.embedding_svc.generate_embedding("test")
            status['connected'] = test is not None and len(test) == 768
            status['embedding_dimensions'] = len(test) if test else 0

            # Test chat LLM
            chat_status = self.chat_svc.health_check()
            status['chat'] = chat_status
        except Exception as e:
            status['error'] = str(e)

        return status


# Instance singleton
ai_service = AltiusAIService()
